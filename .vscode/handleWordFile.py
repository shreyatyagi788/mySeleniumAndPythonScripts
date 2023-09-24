from docx import  Document
import paragraphs
doc = Document()
doc.add_heading("hello world",0) # 0 is for title of the page and 1,2,3...etc are for respective headings.
p=doc.add_paragraph("this is a line of a simple paragraph")
p.add_run(" This is bold text.").bold =True
p.add_run(" This is italic text.").italic=True
doc.add_paragraph("This is item one",style="List Bullet")
doc.add_paragraph("This is item two",style="List Bullet")
doc.add_paragraph("This is item three",style="List Bullet")
doc.add_paragraph("This is item four",style="List Bullet")
doc.add_paragraph("This is item five",style="List Bullet")

doc.save("firstWord.docx")